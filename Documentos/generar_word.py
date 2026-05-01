"""
Convierte librerias_proyecto.md → librerias_proyecto.docx
con estilos aplicados según la sintaxis Markdown.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

MD_FILE = os.path.join(os.path.dirname(__file__), "ARQUITECTURA_SEGURIDAD_RDWATCH.md")
OUT_FILE = os.path.join(os.path.dirname(__file__), "ARQUITECTURA_SEGURIDAD_RDWATCH.docx")

# ─── Colores RD Watch ───────────────────────────────────────────────
COLOR_GOLD   = RGBColor(0xAF, 0x94, 0x4F)   # #AF944F  (dorado marca)
COLOR_DARK   = RGBColor(0x0D, 0x0D, 0x0D)   # #0D0D0D  (casi negro)
COLOR_GRAY   = RGBColor(0x44, 0x44, 0x44)   # #444444  (gris texto)
COLOR_CODE_BG= RGBColor(0xF0, 0xF0, 0xF0)   # fondo código
COLOR_CODE_FG= RGBColor(0x1A, 0x1A, 0x2E)   # texto código

doc = Document()

# ─── Márgenes ───────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ─── Estilos base ───────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = COLOR_GRAY


def set_cell_bg(cell, hex_color):
    """Pone color de fondo a una celda de tabla (para bloques de código)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_DARK
        # Línea dorada decorativa debajo del H1
        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'AF944F')
        border.append(bottom)
        p._p.get_or_add_pPr().append(border)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_GOLD
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK


def add_code_block(lines):
    """Bloque de código en tabla de 1 celda con fondo gris."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F0F0F0')
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = COLOR_CODE_FG
    doc.add_paragraph()  # espacio después del bloque


def apply_inline(paragraph, text):
    """
    Parsea el texto con backtick (`code`) y **negrita**,
    agrega los runs al párrafo con formato correcto.
    """
    # Unificamos los patrones: **bold** y `code`
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)  # rojo suave
        else:
            paragraph.add_run(part)


def add_table_row(cells_text, is_header=False):
    return cells_text  # lo acumulamos y lo procesamos al final


# ─── Parsear el Markdown ─────────────────────────────────────────────
with open(MD_FILE, encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_lines = []
table_rows = []

i = 0
while i < len(lines):
    raw = lines[i].rstrip('\n')

    # ── Bloque de código ─────────────────────────────────────────
    if raw.strip().startswith('```'):
        if not in_code:
            in_code   = True
            code_lines = []
        else:
            in_code = False
            add_code_block(code_lines)
        i += 1
        continue

    if in_code:
        code_lines.append(raw)
        i += 1
        continue

    # ── Línea horizontal ─────────────────────────────────────────
    if raw.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'AF944F')
        border.append(bottom)
        p._p.get_or_add_pPr().append(border)
        i += 1
        continue

    # ── Tabla Markdown ────────────────────────────────────────────
    if raw.strip().startswith('|'):
        # Detección de fila separadora (|---|---|)
        if re.match(r'^\s*\|[\s\-|]+\|\s*$', raw):
            i += 1
            continue
        cells = [c.strip() for c in raw.strip().strip('|').split('|')]
        table_rows.append(cells)
        i += 1
        # Mira si la siguiente es separador o continúa tabla
        if i >= len(lines) or not lines[i].strip().startswith('|'):
            # Renderizar la tabla acumulada
            if table_rows:
                t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                t.style = 'Table Grid'
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = t.cell(r_idx, c_idx)
                        cell.paragraphs[0].clear()
                        p = cell.paragraphs[0]
                        if r_idx == 0:
                            set_cell_bg(cell, 'AF944F')
                            run = p.add_run(cell_text)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            apply_inline(p, cell_text)
                table_rows = []
                doc.add_paragraph()
        continue

    # ── Headings ─────────────────────────────────────────────────
    m = re.match(r'^(#{1,3})\s+(.*)', raw)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        add_heading(text, level)
        i += 1
        continue

    # ── Listas (- o *) ───────────────────────────────────────────
    m = re.match(r'^(\s*)[-*]\s+(.*)', raw)
    if m:
        indent = len(m.group(1)) // 2
        text   = m.group(2)
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent   = Inches(0.3 * (indent + 1))
        p.paragraph_format.space_before  = Pt(2)
        p.paragraph_format.space_after   = Pt(2)
        apply_inline(p, text)
        i += 1
        continue

    # ── Línea vacía ──────────────────────────────────────────────
    if raw.strip() == '':
        doc.add_paragraph()
        i += 1
        continue

    # ── Párrafo normal ───────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    apply_inline(p, raw.strip())
    i += 1

# ─── Guardar ─────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"✅ Documento Word generado: {OUT_FILE}")
